from http.server import BaseHTTPRequestHandler
import json
import tempfile
import os
from pathlib import Path
import PyPDF2
from PIL import Image
from docx import Document
import base64
import io

class handler(BaseHTTPRequestHandler):
    def do_POST(self):
        try:
            # Get content length
            content_length = int(self.headers['Content-Length'])
            
            # Read the body
            body = self.rfile.read(content_length)
            
            # Parse JSON
            data = json.loads(body.decode('utf-8'))
            
            # Get file data and type
            file_data = data.get('file')
            file_type = data.get('type', 'pdf')
            
            if not file_data:
                self.send_error(400, "No file data provided")
                return
            
            # Decode base64 file
            file_bytes = base64.b64decode(file_data.split(',')[1] if ',' in file_data else file_data)
            
            # Extract text based on type
            text = ""
            
            if file_type == 'pdf':
                text = self.extract_from_pdf(file_bytes)
            elif file_type == 'docx':
                text = self.extract_from_docx(file_bytes)
            elif file_type in ['image', 'png', 'jpg', 'jpeg']:
                text = "Image OCR is not available on Vercel. Please use the local Python server for image OCR."
            else:
                self.send_error(400, f"Unsupported file type: {file_type}")
                return
            
            # Send response
            self.send_response(200)
            self.send_header('Content-type', 'application/json')
            self.send_header('Access-Control-Allow-Origin', '*')
            self.end_headers()
            
            response = {
                'text': text,
                'length': len(text)
            }
            
            self.wfile.write(json.dumps(response).encode('utf-8'))
            
        except Exception as e:
            self.send_response(500)
            self.send_header('Content-type', 'application/json')
            self.send_header('Access-Control-Allow-Origin', '*')
            self.end_headers()
            
            error_response = {
                'error': str(e)
            }
            
            self.wfile.write(json.dumps(error_response).encode('utf-8'))
    
    def do_OPTIONS(self):
        self.send_response(200)
        self.send_header('Access-Control-Allow-Origin', '*')
        self.send_header('Access-Control-Allow-Methods', 'POST, OPTIONS')
        self.send_header('Access-Control-Allow-Headers', 'Content-Type')
        self.end_headers()
    
    def extract_from_pdf(self, file_bytes):
        """Extract text from PDF"""
        try:
            pdf_file = io.BytesIO(file_bytes)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text = ""
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            
            if text.strip():
                return text.strip()
            else:
                return "No text found in PDF. It may be a scanned document requiring OCR."
                
        except Exception as e:
            raise Exception(f"PDF extraction failed: {str(e)}")
    
    def extract_from_docx(self, file_bytes):
        """Extract text from Word document"""
        try:
            docx_file = io.BytesIO(file_bytes)
            doc = Document(docx_file)
            
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return text.strip()
            
        except Exception as e:
            raise Exception(f"Word document extraction failed: {str(e)}")

