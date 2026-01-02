"""
Flask API server for OCR with Tesseract support
Runs in Docker container
"""
from flask import Flask, request, jsonify
from flask_cors import CORS
import os
import tempfile
from pathlib import Path
import logging
import PyPDF2
from PIL import Image
import pytesseract
from pdf2image import convert_from_path
from docx import Document

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = Flask(__name__)
CORS(app)  # Enable CORS for all routes

class OCRProcessor:
    """OCR processing class"""
    
    def extract_from_pdf(self, file_path: Path) -> str:
        """Extract text from PDF using PyPDF2 and OCR fallback"""
        logger.info(f"Extracting text from PDF: {file_path}")
        
        try:
            # Try text extraction first
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                
                if text.strip():
                    logger.info(f"Text extraction successful: {len(text)} characters")
                    return text.strip()
        except Exception as e:
            logger.warning(f"Text extraction failed: {e}, trying OCR")
        
        # Fallback to OCR for scanned PDFs
        try:
            logger.info("Converting PDF to images for OCR")
            images = convert_from_path(str(file_path), dpi=300)
            text = ""
            for i, image in enumerate(images):
                logger.info(f"Processing page {i+1}/{len(images)}")
                page_text = pytesseract.image_to_string(image)
                text += page_text + "\n"
            
            logger.info(f"OCR extraction successful: {len(text)} characters")
            return text.strip()
        except Exception as e:
            raise Exception(f"Failed to extract text from PDF: {str(e)}")
    
    def extract_from_image(self, file_path: Path) -> str:
        """Extract text from image using pytesseract"""
        logger.info(f"Extracting text from image: {file_path}")
        
        try:
            image = Image.open(file_path)
            text = pytesseract.image_to_string(image)
            logger.info(f"OCR extracted {len(text)} characters")
            return text.strip()
        except Exception as e:
            raise Exception(f"Failed to perform OCR on image: {str(e)}")
    
    def extract_from_docx(self, file_path: Path) -> str:
        """Extract text from Word document"""
        logger.info(f"Extracting text from Word document: {file_path}")
        
        try:
            doc = Document(file_path)
            text = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            logger.info(f"Extracted {len(text)} characters from Word document")
            return text.strip()
        except Exception as e:
            raise Exception(f"Failed to extract text from Word document: {str(e)}")

processor = OCRProcessor()

@app.route('/health', methods=['GET'])
def health():
    """Health check endpoint"""
    return jsonify({'status': 'healthy', 'tesseract': 'available'}), 200

@app.route('/api/ocr', methods=['POST', 'OPTIONS'])
def ocr():
    """OCR endpoint"""
    if request.method == 'OPTIONS':
        return '', 200
    
    try:
        # Check if file is in request
        if 'file' not in request.files:
            return jsonify({'error': 'No file provided'}), 400
        
        file = request.files['file']
        
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400
        
        # Get file extension
        filename = file.filename.lower()
        ext = filename.split('.')[-1] if '.' in filename else ''
        
        # Save file temporarily
        with tempfile.NamedTemporaryFile(delete=False, suffix=f'.{ext}') as tmp_file:
            file.save(tmp_file.name)
            tmp_path = Path(tmp_file.name)
        
        try:
            # Extract text based on file type
            if ext == 'pdf':
                text = processor.extract_from_pdf(tmp_path)
            elif ext == 'docx':
                text = processor.extract_from_docx(tmp_path)
            elif ext in ['jpg', 'jpeg', 'png', 'gif', 'bmp', 'tiff']:
                text = processor.extract_from_image(tmp_path)
            else:
                return jsonify({'error': f'Unsupported file type: {ext}'}), 400
            
            return jsonify({
                'text': text,
                'length': len(text),
                'filename': file.filename
            }), 200
            
        finally:
            # Clean up temp file
            if tmp_path.exists():
                tmp_path.unlink()
    
    except Exception as e:
        logger.error(f"OCR error: {str(e)}")
        return jsonify({'error': str(e)}), 500

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 8000))
    logger.info(f"Starting OCR API server on port {port}")
    app.run(host='0.0.0.0', port=port, debug=False)

